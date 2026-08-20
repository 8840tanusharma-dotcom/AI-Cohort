'''Extract and normalize unstructured text from PDFs, DOCX, OCR scans, and Web FAQs.'''

from pathlib import Path
import re
from typing import Optional

ROOT = Path(__file__).resolve().parents[1]
RAW_TEXT_DIR = ROOT / "raw_text"


def normalize_text(text: str) -> str:
    '''Normalize extracted text by cleaning whitespaces, headers/footers, and encoding artifacts.'''
    if not text:
        return ""
    # Normalize unicode / line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Remove excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip trailing whitespace on each line
    lines = [line.rstrip() for line in text.split('\n')]
    # Reassemble and strip overall
    cleaned = '\n'.join(lines).strip()
    return cleaned


def extract_pdf_text(pdf_path: Path) -> str:
    '''Extract structured text from a PDF document using pdfplumber.'''
    import pdfplumber

    extracted_pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            extracted_pages.append(page_text)
    return normalize_text("\n\n".join(extracted_pages))


def extract_docx_text(docx_path: Path) -> str:
    '''Extract structured text from a Word document using python-docx.'''
    import docx

    doc = docx.Document(docx_path)
    full_text = [p.text for p in doc.paragraphs if p.text.strip()]
    # Extract any table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)
    return normalize_text("\n\n".join(full_text))


def extract_ocr_text(image_or_pdf_path: Path, tesseract_cmd: Optional[str] = None) -> str:
    '''Extract text from a scanned document image or scanned PDF using pytesseract.'''
    import pytesseract
    from PIL import Image

    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    if image_or_pdf_path.suffix.lower() == ".pdf":
        from pdf2image import convert_from_path
        images = convert_from_path(str(image_or_pdf_path))
        texts = [pytesseract.image_to_string(img) for img in images]
        return normalize_text("\n\n".join(texts))
    else:
        image = Image.open(image_or_pdf_path)
        return normalize_text(pytesseract.image_to_string(image))


def scrape_faq_page(url: str) -> str:
    '''Scrape public FAQ or provider network documentation using requests and BeautifulSoup.'''
    import requests
    from bs4 import BeautifulSoup

    headers = {"User-Agent": "HealthCoverageBot/1.0 (Public Documentation Ingestion)"}
    response = requests.get(url, headers=headers, timeout=10)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove script and style elements
    for script in soup(["script", "style", "nav", "footer"]):
        script.extract()

    # Extract text from main content
    main_content = soup.find("main") or soup.find("article") or soup.find("body")
    raw_text = main_content.get_text(separator="\n") if main_content else soup.get_text(separator="\n")
    return normalize_text(raw_text)


if __name__ == "__main__":
    RAW_TEXT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Unstructured extraction module ready. Target directory: {RAW_TEXT_DIR}")
    for file in RAW_TEXT_DIR.glob("*.txt"):
        print(f" - {file.name}: {file.stat().st_size} bytes")
